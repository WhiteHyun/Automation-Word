def auto_process(student_folder_name: str, result_folder_name: str = "result", verbose: bool = False):
    """
    학생들의 중간고사 문제지를 자동으로 처리해주는 함수입니다.

    Parameter
    ---------
    student_folder_name: str
        학생들의 코드파일을 하나의 폴더에 전부 넣어두셔야 합니다.
        이 때 해당 폴더의 제목을 파라미터로 받습니다.

    result_folder_name: str
        자동화된 결과물을 저장할 폴더명입니다. 기본값은 `result` 입니다.

    verbose: bool
        무엇을 처리하는지 로그를 보여줍니다.

    Raises
    ------
    AssertionError
        다음과 같은 경우 해당 오류가 발생합니다.

        1. 입력받은 폴더파일에 학생들의 코드파일이 존재하지 않을경우

        2. 결과물을 저장할 폴더가 이미 생성되어져 있는 경우
            - 본래 존재하는 폴더일 경우 파일에 손상을 끼칠 수 있기 때문에 따로 처리해두었습니다.

    Example
    -------
    >>> auto_process("중간고사A반")
    # 중간고사A반 폴더에 있는 학생들의 코드파일을 전부 처리하여 result 폴더 내에 정리해줌
    """

    import os
    import shutil
    from docx import Document

    # 학생들의 코드 파일
    file_list = os.listdir(student_folder_name)
    assert file_list, "폴더에 아무런 파일도 없습니다!"

    # 폴더가 만들어져있지 않을 경우
    if not os.path.exists(result_folder_name):
        os.makedirs(result_folder_name)  # 폴더 생성

    if verbose:
        print("Starting...")

    # 각 학생들의 파일을 읽어옴
    for file in file_list:
        if file[-5:] != ".docx":  # docx 파일이 아니면 슥 이동
            continue

        # 학번 인식
        school_number = file[:-5].split("file_")[-1]
        new_file_folder = f"{result_folder_name}/{school_number}"

        # 학생 폴더 생성
        if os.path.exists(new_file_folder):
            shutil.rmtree(new_file_folder)  # 기존 폴더 삭제
        os.makedirs(new_file_folder)
        if verbose:
            print(f"{new_file_folder} create..")

        # 코드파일 읽음 (docx 파일)
        doc = Document(f"{student_folder_name}/{file}")
        if verbose:
            print(f"{file} read..")

        problem_list = list(map(lambda table: True if len(
            table.rows) & 1 == 0 else False, doc.tables))

        # 파일 검사
        assert problem_list.count(True) == len(
            problem_list), f"{problem_list.index(False)+1}번 문제의 양식이 잘못되었습니다."

        number = 0
        for table in doc.tables:  # 각 문제들에 대해 작업
            file_name = ""
            number += 1
            i = 0
            if verbose:
                print(f"{school_number} problem {number} start")

            # 학생의 각 문제에 따른 폴더 생성
            os.makedirs(f"{new_file_folder}/problem{number}")

            for row in table.rows:
                if i & 1 == 0:  # 만약 짝수인 경우 해당 코드의 파일명임, ex) activity_main.xml, MainActivity.kt, ..
                    file_name = row.cells[0].paragraphs[0].text

                else:  # 문제에 대한 코드인 경우
                    if verbose:
                        print(
                            f"{school_number} problem {number}, {file_name} process..")
                    with open(f"./{new_file_folder}/problem{number}/{file_name}", "w", encoding="UTF-8") as f:
                        f.writelines("\n".join(
                            list(map(lambda paragraph: paragraph.text, row.cells[0].paragraphs))))
                i += 1
            if verbose:
                print(f"{school_number} problem {number} finished")
        if verbose:
            print(f"{file} done..")
    if verbose:
        print("All Done.")


if __name__ == "__main__":
    auto_process("중간고사B", "mid-term B반", verbose=True)
