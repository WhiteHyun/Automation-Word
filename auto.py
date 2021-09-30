import pandas as pd
from docx import Document
from docx.shared import Cm

MONTH = 0
DAY = 1

global_time = {
    "글로벌소양과문화컨텐츠": ("1300", "1500"),
    "문화예술로배우는생명의기원": ("1030", "1330"),
    "바이러스학1": ("0900", "1030"),
    "바이러스학2": ("0900", "1030"),
    "정보보호론1": ("1330", "1500"),
    "정보보호론2": ("1330", "1500"),
    "지능정보시스템": ("1930", "2230"),
    "창업마케팅AtoZ": ("1200", "1500")
}


def make_work_log(date: tuple, time: tuple, content: str, image_url: str, lecture: str):
    """

    Example
    -------
    >>> make_work_log(date=("9", "7"),
                      time=("0900", "1030"),
                      content="안녕하세요\\n만나서반갑습니다.",
                      image_url="test_image.jpg"
                      lecture="지능정보시스템")
    작성된 문서가 저장

    """
    doc = Document("양식.docx")
    tables = doc.tables

    work_day = tables[0].rows[3].cells[1].paragraphs[0]
    work_times = tables[0].rows[3].cells[3].paragraphs

    work_content = tables[0].rows[5].cells[1]
    work_picture = tables[0].rows[6].cells[1].paragraphs[0].add_run()

    # 근로일시 작성
    work_day.text = f"21년   {date[MONTH]}월  {date[DAY]}일"

    # 근로시간 작성
    for i, work_time in zip((0, 1), work_times):
        work_time.text = f"  {time[i][:2]}시  {time[i][2:]}분 {work_time.text[11:13]}"

    # 근로내용 작성
    for sentence in content.split("\n"):
        work_content.add_paragraph(sentence)

    # 근로사진 첨부

    work_picture.add_picture(image_url, width=Cm(11.62), height=Cm(7.24))

    doc.save(f"note/{date[MONTH]}월{date[DAY]}일 - {lecture}.docx")


def preprocessing(series):

    # 날짜 전처리
    date_length = str(series["날짜"])
    if date_length == 4:  # 10월 이후
        month = int(date_length[:2])
        day = int(date_length[2:])
    else:
        month = int(date_length[:1])
        day = int(date_length[1:])
    date = (month, day)

    # 시간 전처리
    time = global_time[series["강의명"]]

    # 내용 전처리
    text = series["내용"]
    # csv에서 개행을 #!#로 대체해서 넣어놓았기 때문에 다시 바꿔줌
    text = text.replace("#!#", "\n")

    # 이미지 주소 전처리
    image_url = f"image/{series['이미지']}.png"
    return date, time, text, image_url


if __name__ == "__main__":
    df = pd.read_csv("word.csv")

    # csv의 과목명과 global_time (dict)의 길이가 무조건 같아야 함
    assert(len(set(df.loc[:, "강의명"])) == len(global_time))
    for i in df.index:
        date, time, content, image_url = preprocessing(df.iloc[i, :])
        make_work_log(date=date,
                      time=time,
                      content=content,
                      image_url=image_url,
                      lecture=df.loc[i, "강의명"])
