from docx import Document
from docx.shared import Cm
import pandas as pd
from class_time import class_datetime_dict

MONTH = 0
DAY = 1


class DailyLog:
    """
    근로일지를 작성하기 위해 사용되는 클래스입니다.
    각 프로퍼티를 가지고 근로일지를 작성해줍니다.
    """

    def __init__(
        self, date: tuple, time: tuple, content: str, image_link: str, lecture: str
    ) -> None:
        self.date = date
        self.time = time
        self.content = content
        self.image = image_link
        self.lecture = lecture

    def make_work_log(self):
        """

        Example
        -------
        >>> make_work_log(date=("9", "7"),
                        time=("0900", "1030"),
                        content="안녕하세요\\n만나서반갑습니다.",
                        image_url="test_image.jpg"
                        lecture="지능정보시스템")
        작성된 문서가 저장

        """
        doc = Document("양식.docx")
        tables = doc.tables

        work_day = tables[0].rows[3].cells[1].paragraphs[0]
        work_times = tables[0].rows[3].cells[3].paragraphs

        work_content = tables[0].rows[5].cells[1]
        work_picture = tables[0].rows[6].cells[1].paragraphs[0].add_run()

        # 근로일시 작성
        work_day.text = f"21년   {self.date[MONTH]}월  {self.date[DAY]}일"

        # 근로시간 작성
        for i, work_time in zip((0, 1), work_times):
            work_time.text = (
                f"  {self.time[i][:2]}시  {self.time[i][2:]}분 {work_time.text[11:13]}"
            )

        # 근로내용 작성
        for sentence in self.content.split("\n"):
            work_content.add_paragraph(sentence)

        # 근로사진 첨부

        work_picture.add_picture(
            self.image, width=Cm(11.62), height=Cm(7.24))

        save_url = f"note/{self.date[MONTH]}월{self.date[DAY]}일 - {self.lecture}.docx"
        print(f"saving {save_url[5:]}...")

        doc.save(
            f"note/{self.date[MONTH]}월{self.date[DAY]}일 - {self.lecture}.docx")


def preprocessing(
    series: pd.Series,
) -> tuple[tuple[int, int], tuple[str, str], str, str]:

    # 날짜 전처리
    date_length = str(series["날짜"])
    assert len(date_length) == 4, f"날짜 형식이 잘못되었습니다. {date_length=}"
    month = int(date_length[:2])
    day = int(date_length[2:])
    date = (month, day)

    # 시간 전처리
    time = class_datetime_dict[series["강의명"]][1:]

    # 내용 전처리
    text = series["내용"]
    # csv에서 개행을 #!#로 대체해서 넣어놓았기 때문에 다시 바꿔줌
    text = text.replace("#!#", "\n")

    # 이미지 주소 전처리
    image_url = f"image/{series['이미지']}.png"
    return date, time, text, image_url
